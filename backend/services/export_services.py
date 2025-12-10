# services/export_service.py
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph

class PDFExporter:
    def export(self, documentation, output_path):
        doc = SimpleDocTemplate(output_path)
        story = []
        styles = getSampleStyleSheet()
        
        # Add title
        story.append(Paragraph(f"# {documentation.title}", styles['Heading1']))
        
        # Add code blocks with syntax highlighting
        for block in documentation.code_blocks:
            story.append(Paragraph(f"```{block.language}", styles['Code']))
            story.append(Paragraph(block.code, styles['Code']))
            story.append(Paragraph("```", styles['Code']))
        
        doc.build(story)

# services/export_service.py
from docx import Document
from docx.shared import Inches

class DOCXExporter:
    def export(self, documentation, output_path):
        doc = Document()
        
        # Add title
        doc.add_heading(documentation.title, 0)
        
        # Add code blocks
        for block in documentation.code_blocks:
            doc.add_paragraph(f"Language: {block.language}")
            doc.add_paragraph(block.code, style='Code')
            
        doc.save(output_path)