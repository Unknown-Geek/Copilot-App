import re
import json
import logging
import datetime
from io import BytesIO
from typing import List, Dict, Any, Literal, Optional, Union
from dataclasses import dataclass, field

# Remove pdfkit import
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from pygments import lex
from pygments.lexers import get_lexer_by_name

try:
    import radon.metrics
    import radon.complexity
    from radon.visitors import ComplexityVisitor
    RADON_AVAILABLE = True
except ImportError:
    RADON_AVAILABLE = False
    logging.warning("radon package not found. Advanced metrics will be limited.")

@dataclass
class CodeBlock:
    content: str
    language: str
    line_number: int

@dataclass
class Documentation:
    """Documentation model class"""
    title: str 
    description: str 
    code_blocks: List[CodeBlock]
    language: str
    generated_at: str
    metrics: Dict[str, Any] = field(default_factory=dict)

    def __init__(self, title: str, description: str, language: str, code_blocks: List[CodeBlock], metrics: Dict[str, Any]):
        self.title = title
        self.description = description
        self.language = language
        self.code_blocks = code_blocks
        self.metrics = metrics
        self.generated_at = datetime.datetime.now().isoformat()

    def to_dict(self) -> Dict[str, Any]:
        """Convert Documentation to dictionary"""
        return {
            'title': self.title,
            'description': self.description,
            'code_blocks': [
                {
                    'content': block.content,
                    'language': block.language,
                    'line_number': block.line_number
                } for block in self.code_blocks
            ],
            'language': self.language,
            'generated_at': self.generated_at,
            'metrics': self.metrics
        }

class DocumentationGenerator:
    def __init__(self):
        # Add RADON_AVAILABLE as class attribute
        self.RADON_AVAILABLE = RADON_AVAILABLE
        self.exporters = {
            'markdown': self._export_markdown,
            'html': self._export_html,
            'json': self._export_json,
            'pdf': self._export_pdf,
            'docx': self._export_docx
        }
        self.supported_languages = {
            'python': self._parse_python,
            'javascript': self._parse_javascript,
            'typescript': self._parse_javascript,
            'java': self._parse_java,
            'cpp': self._parse_cpp,
            'csharp': self._parse_csharp
        }
        self.templates = {
            'default': {
                'title': '# {title}',
                'description': '{description}',
                'code_blocks': '## Code\n```{language}\n{content}\n```',
                'metrics': '## Metrics\n{metrics}'
            },
            'detailed': {
                'title': '# {title}\n\n_Generated at: {generated_at}_',
                'description': '## Overview\n\n{description}',
                'code_blocks': '## Code Blocks\n\n### Block at line {line_number}\n```{language}\n{content}\n```\n\n**Analysis:**\n- Lines of code: {loc}\n- Complexity: {complexity}',
                'metrics': '## Project Metrics\n\n```json\n{metrics}\n```',
                'toc': '## Table of Contents\n\n{toc}'
            }
        }
        self.supported_formats = {
            'markdown': self._export_markdown,
            'html': self._export_html,
            'json': self._export_json
        }
        self.logger = logging.getLogger(__name__)

    def _calculate_metrics(self, code_blocks: List[CodeBlock]) -> Dict[str, Any]:
        total_blocks = len(code_blocks)
        total_lines = sum(block.content.count('\n') + 1 for block in code_blocks)
        average_block_size = total_lines / total_blocks if total_blocks > 0 else 0
        average_complexity = 1.0  # Placeholder for actual complexity calculation
        max_complexity = 1.0  # Placeholder for actual complexity calculation
        
        return {
            'total_blocks': total_blocks,
            'total_lines': total_lines,
            'average_block_size': average_block_size,
            'average_complexity': average_complexity,
            'max_complexity': max_complexity
        }

    def _parse_code_blocks(self, code: str, language: str) -> List[CodeBlock]:
        """Parse code into code blocks based on the language."""
        if language == 'python':
            return self._parse_python(code)
        elif language == 'javascript':
            return self._parse_javascript(code)
        elif language == 'java':
            return self._parse_java(code)
        elif language == 'cpp':
            return self._parse_cpp(code)
        elif language == 'csharp':
            return self._parse_csharp(code)
        else:
            raise ValueError(f"Unsupported language: {language}")

    def _parse_python(self, code: str) -> List[CodeBlock]:
        """Parse Python code into code blocks."""
        # Simple example: split by function definitions
        blocks = []
        lines = code.split('\n')
        current_block = []
        for line in lines:
            if line.strip().startswith('def ') or line.strip().startswith('class '):
                if current_block:
                    blocks.append(CodeBlock(content='\n'.join(current_block)))
                    current_block = []
            current_block.append(line)
        if current_block:
            blocks.append(CodeBlock(content='\n'.join(current_block)))
        return blocks

    def _parse_javascript(self, code: str) -> List[CodeBlock]:
        """Parse JavaScript code into code blocks."""
        # Simple example: split by function definitions
        blocks = []
        lines = code.split('\n')
        current_block = []
        for line in lines:
            if line.strip().startswith('function ') or line.strip().startswith('class '):
                if current_block:
                    blocks.append(CodeBlock(content='\n'.join(current_block)))
                    current_block = []
            current_block.append(line)
        if current_block:
            blocks.append(CodeBlock(content='\n'.join(current_block)))
        return blocks

    def _parse_java(self, code: str) -> List[CodeBlock]:
        """Parse Java code into code blocks."""
        # Simple example: split by class and method definitions
        blocks = []
        lines = code.split('\n')
        current_block = []
        for line in lines:
            if line.strip().startswith('public class ') or line.strip().startswith('public void '):
                if current_block:
                    blocks.append(CodeBlock(content='\n'.join(current_block)))
                    current_block = []
            current_block.append(line)
        if current_block:
            blocks.append(CodeBlock(content='\n'.join(current_block)))
        return blocks

    def _parse_cpp(self, code: str) -> List[CodeBlock]:
        """Parse C++ code into code blocks."""
        blocks = []
        lines = code.split('\n')
        current_block = []
        for line in lines:
            if line.strip().startswith('void ') or line.strip().startswith('class '):
                if current_block:
                    blocks.append(CodeBlock(content='\n'.join(current_block)))
                    current_block = []
            current_block.append(line)
        if current_block:
            blocks.append(CodeBlock(content='\n'.join(current_block)))
        return blocks

    def _parse_csharp(self, code: str) -> List[CodeBlock]:
        """Parse C# code into code blocks."""
        blocks = []
        lines = code.split('\n')
        current_block = []
        for line in lines:
            if line.strip().startswith('public void ') or line.strip().startswith('public class '):
                if current_block:
                    blocks.append(CodeBlock(content='\n'.join(current_block)))
                    current_block = []
            current_block.append(line)
        if current_block:
            blocks.append(CodeBlock(content='\n'.join(current_block)))
        return blocks

    def generate(self, code: str, language: str, title: Optional[str] = None, 
                 description: Optional[str] = None) -> Documentation:
        """
        Generate documentation for the given source code.

        Args:
            code (str): The source code to generate documentation for
            language (str): Programming language of the source code
            title (str, optional): Custom title for the documentation
            description (str, optional): Custom description for the documentation

        Returns:
            Documentation: Generated documentation object

        Raises:
            ValueError: If code is empty or language is not supported
        """
        if not code:
            raise ValueError("Code cannot be empty")
        if language not in self.supported_languages:
            raise ValueError(f"Unsupported language: {language}")

        code_blocks = self._parse_code_blocks(code, language)
        metrics = self._calculate_metrics(code_blocks)

        return Documentation(
            title=title or f"{language.capitalize()} Documentation",
            description=description or "Auto-generated documentation",
            language=language,
            code_blocks=code_blocks,
            metrics=metrics
        )

    def export_documentation(self, doc: Documentation, format: str = 'markdown',
                        template: str = 'default', output_path: str = None) -> Union[str, bytes]:
        """Export documentation to the specified format."""
        if template not in self.templates:
            raise ValueError(f"Invalid template: {template}")
            
        if format not in self.exporters:
            raise ValueError(f"Invalid format: {format}")
        
        content = self.exporters[format](doc, template)
        
        if output_path:
            mode = 'wb' if isinstance(content, bytes) else 'w'
            with open(output_path, mode) as f:
                f.write(content)
        
        return content

    def _parse_python(self, code: str) -> List[CodeBlock]:
        """Parse Python code into code blocks."""
        blocks = []
        lines = code.split('\n')
        current_block = []
        line_number = 1

        for line in lines:
            stripped = line.strip()
            # Detect function or class definitions
            if stripped.startswith('def ') or stripped.startswith('class '):
                if current_block:
                    blocks.append(CodeBlock(
                        content='\n'.join(current_block),
                        language='python',
                        line_number=line_number - len(current_block)
                    ))
                current_block = [line]
            elif current_block:
                current_block.append(line)
            line_number += 1

        # Add the last block
        if current_block:
            blocks.append(CodeBlock(
                content='\n'.join(current_block),
                language='python',
                line_number=line_number - len(current_block)
            ))

        return blocks

    def _parse_javascript(self, code: str) -> List[CodeBlock]:
        """Parse JavaScript/TypeScript code into code blocks."""
        blocks = []
        lines = code.split('\n')
        current_block = []
        line_number = 1

        for line in lines:
            stripped = line.strip()
            # Detect function or class definitions
            if (stripped.startswith('function ') or 
                stripped.startswith('class ') or 
                stripped.startswith('const ') and '=> {' in stripped):
                if current_block:
                    blocks.append(CodeBlock(
                        content='\n'.join(current_block),
                        language='javascript',
                        line_number=line_number - len(current_block)
                    ))
                current_block = [line]
            elif current_block:
                current_block.append(line)
            line_number += 1

        # Add the last block
        if current_block:
            blocks.append(CodeBlock(
                content='\n'.join(current_block),
                language='javascript',
                line_number=line_number - len(current_block)
            ))

        return blocks

    def _parse_java(self, source_code: str) -> List[CodeBlock]:
        """Parse Java source code and extract code blocks."""
        code_blocks = []
        lines = source_code.split('\n')
        
        current_block = []
        in_comment = False
        
        for i, line in enumerate(lines, 1):
            stripped_line = line.strip()
            
            # Handle multi-line comments
            if '/*' in stripped_line:
                in_comment = True
            if '*/' in stripped_line:
                in_comment = False
                continue
                
            # Skip comments
            if stripped_line.startswith('//') or in_comment:
                continue
                
            # Check for class, interface, or method declarations
            if re.match(r'^\s*(public|private|protected)?\s*(class|interface|enum|@interface|abstract\s+class|\w+\s+\w+\s*\()', line):
                if current_block:
                    code_blocks.append(CodeBlock(
                        content='\n'.join(current_block),
                        language='java',
                        line_number=i - len(current_block)
                    ))
                current_block = [line]
            elif current_block:
                current_block.append(line)
                
        if current_block:
            code_blocks.append(CodeBlock(
                content='\n'.join(current_block),
                language='java',
                line_number=len(lines) - len(current_block) + 1
            ))
            
        return code_blocks

    def _parse_cpp(self, source_code: str) -> List[CodeBlock]:
        """Parse C++ source code and extract code blocks."""
        code_blocks = []
        lines = source_code.split('\n')
        
        current_block = []
        in_comment = False
        template_depth = 0
        
        for i, line in enumerate(lines, 1):
            stripped_line = line.strip()
            
            # Handle multi-line comments
            if '/*' in stripped_line:
                in_comment = True
            if '*/' in stripped_line:
                in_comment = False
                continue
                
            # Skip comments
            if stripped_line.startswith('//') or in_comment:
                continue
                
            # Track template depth
            template_depth += stripped_line.count('<') - stripped_line.count('>')
            
            # Check for function, class, or struct declarations
            if template_depth == 0 and re.match(r'^\s*(template\s*<.*>)?\s*(class|struct|enum|union|\w+\s+\w+\s*\()', line):
                if current_block:
                    code_blocks.append(CodeBlock(
                        content='\n'.join(current_block),
                        language='cpp',
                        line_number=i - len(current_block)
                    ))
                current_block = [line]
            elif current_block:
                current_block.append(line)
                
        if current_block:
            code_blocks.append(CodeBlock(
                content='\n'.join(current_block),
                language='cpp',
                line_number=len(lines) - len(current_block) + 1
            ))
            
        return code_blocks

    def _parse_csharp(self, source_code: str) -> List[CodeBlock]:
        """Parse C# source code and extract code blocks."""
        code_blocks = []
        lines = source_code.split('\n')
        
        current_block = []
        in_comment = False
        
        for i, line in enumerate(lines, 1):
            stripped_line = line.strip()
            
            # Handle multi-line comments
            if '/*' in stripped_line:
                in_comment = True
            if '*/' in stripped_line:
                in_comment = False
                continue
                
            # Skip comments
            if stripped_line.startswith('//') or in_comment:
                continue
                
            # Check for class, interface, method, or property declarations
            if re.match(r'^\s*(public|private|protected|internal)?\s*(class|interface|enum|struct|record|async\s+Task|\w+\s+\w+\s*\(|\w+\s+\w+\s*{)', line):
                if current_block:
                    code_blocks.append(CodeBlock(
                        content='\n'.join(current_block),
                        language='csharp',
                        line_number=i - len(current_block)
                    ))
                current_block = [line]
            elif current_block:
                current_block.append(line)
                
        if current_block:
            code_blocks.append(CodeBlock(
                content='\n'.join(current_block),
                language='csharp',
                line_number=len(lines) - len(current_block) + 1
            ))
            
        return code_blocks

    def _generate_metrics(self, code_blocks: List[CodeBlock]) -> Dict[str, Any]:
        """Generate metrics for the parsed code blocks."""
        total_lines = sum(len(block.content.split('\n')) for block in code_blocks)
        complexities = [self._calculate_complexity(block) for block in code_blocks]
        
        return {
            'total_blocks': len(code_blocks),
            'total_lines': total_lines,
            'average_block_size': total_lines / len(code_blocks) if code_blocks else 0,
            'average_complexity': sum(complexities) / len(complexities) if complexities else 0,
            'max_complexity': max(complexities) if complexities else 0
        }

    def _calculate_complexity(self, block: CodeBlock) -> int:
        """Calculate cyclomatic complexity for a code block."""
        complexity = 1
        code = block.content.lower()
        
        # Language-specific complexity patterns
        patterns = {
            'python': [r'\bif\b', r'\belif\b', r'\bfor\b', r'\bwhile\b', r'\band\b', r'\bor\b', r'\bcatch\b', r'\bwith\b'],
            'javascript': [r'\bif\b', r'\belse\s+if\b', r'\bfor\b', r'\bwhile\b', r'\b\&\&\b', r'\b\|\|\b', r'\bcatch\b', r'\bcase\b'],
            'java': [r'\bif\b', r'\belse\s+if\b', r'\bfor\b', r'\bwhile\b', r'\b\&\&\b', r'\b\|\|\b', r'\bcatch\b', r'\bcase\b'],
            'cpp': [r'\bif\b', r'\belse\s+if\b', r'\bfor\b', r'\bwhile\b', r'\b\&\&\b', r'\b\|\|\b', r'\bcatch\b', r'\bcase\b'],
            'csharp': [r'\bif\b', r'\belse\s+if\b', r'\bfor\b', r'\bwhile\b', r'\b\&\&\b', r'\b\|\|\b', r'\bcatch\b', r'\bcase\b']
        }
        
        # Add complexity for each control flow pattern found
        for pattern in patterns.get(block.language, patterns['python']):
            complexity += len(re.findall(pattern, code))
        
        return complexity

    def _calculate_advanced_metrics(self, code_block: CodeBlock) -> Dict[str, Any]:
        """Calculate advanced metrics for a code block."""
        metrics = {
            'loc': len(code_block.content.split('\n')),
            'sloc': len([l for l in code_block.content.split('\n') if l.strip()]),
            'comments': len(re.findall(r'(?://|#|/\*|\*/|"""|\'\'\')', code_block.content)),
            'complexity': self._calculate_complexity(code_block),
            'token_count': len(list(lex(code_block.content, 
                get_lexer_by_name(code_block.language)))),
            'char_count': len(code_block.content)
        }
        
        # Add advanced metrics only if radon is available
        if self.RADON_AVAILABLE and code_block.language == 'python':
            try:
                # Fix: Ensure code block has proper indentation for radon
                code_content = code_block.content.strip()
                if not code_content.endswith('\n'):
                    code_content += '\n'
                    
                complexity_visitor = ComplexityVisitor.from_code(code_content)
                metrics.update({
                    'cognitive_complexity': complexity_visitor.total_complexity,
                    'functions': len(complexity_visitor.functions),
                    'classes': len(complexity_visitor.classes),
                    'halstead_metrics': radon.metrics.h_visit(code_content),
                    'maintainability_index': radon.metrics.mi_visit(code_content, True)
                })
            except Exception as e:
                logging.warning(f"Failed to calculate advanced metrics: {str(e)}")
                # Add default values if calculation fails
                metrics.update({
                    'cognitive_complexity': 0,
                    'functions': 0,
                    'classes': 0,
                    'halstead_metrics': {},
                    'maintainability_index': 0
                })
                
        return metrics

    def _calculate_python_metrics(self, code_block: CodeBlock) -> Dict[str, Any]:
        """Calculate Python-specific metrics."""
        metrics = {}
        if RADON_AVAILABLE:
            try:
                complexity_visitor = ComplexityVisitor.from_code(code_block.content)
                metrics.update({
                    'cognitive_complexity': complexity_visitor.total_complexity,
                    'functions': len(complexity_visitor.functions),
                    'classes': len(complexity_visitor.classes),
                    'halstead_metrics': radon.metrics.h_visit(code_block.content),
                    'maintainability_index': radon.metrics.mi_visit(code_block.content, True)
                })
            except Exception as e:
                logging.warning(f"Failed to calculate Python metrics: {str(e)}")
        return metrics

    def save_documentation(self, doc: Documentation, output_path: str) -> None:
        """Save the generated documentation to a file."""
        markdown_content = self._export_markdown(doc, template='default')
        with open(output_path, 'w') as file:
            file.write(markdown_content)

    def export_to_markdown(self, doc: Documentation, output_path: str) -> None:
        """Export the documentation to a Markdown file."""
        markdown_content = self._export_markdown(doc, template='default')
        with open(output_path, 'w') as file:
            file.write(markdown_content)

    def _export_markdown(self, doc: Documentation, template: str) -> str:
        """Export documentation to markdown format."""
        template_config = self.templates[template]
        content = f"# {doc.title}\n\n"
        
        if template == 'detailed':
            content += "## Table of Contents\n"
            for i, block in enumerate(doc.code_blocks):
                content += f"{i+1}. Code Block {i+1}\n"
            content += "\n## Code Blocks\n"
        
        for block in doc.code_blocks:
            content += f"```{doc.language}\n{block.content}\n```\n\n"
        
        if template == 'detailed':
            content += "## Project Metrics\n"
            for key, value in doc.metrics.items():
                content += f"- {key}: {value}\n"
        
        return content

    def _export_html(self, doc: Documentation, template: str) -> str:
        """Export documentation to HTML format."""
        template_config = self.templates[template]
        return f"""<!DOCTYPE html>
        <html>
        <body>
        <div class="code-block">
        <pre><code>{doc.code_blocks[0].content}</code></pre>
        </div>
        </body>
        </html>"""

    def _export_json(self, doc: Documentation, template: str) -> str:
        """Export documentation to JSON format."""
        return json.dumps({
            'title': doc.title,
            'description': doc.description,
            'language': doc.language,
            'metrics': doc.metrics,
            'code_blocks': [
                {
                    'content': block.content,
                    'language': block.language,
                    'line_number': block.line_number,
                    'metrics': {
                        'loc': len(block.content.split('\n')),
                        'complexity': self._calculate_complexity(block)
                    }
                }
                for block in doc.code_blocks
            ]
        }, indent=2)

    def _export_pdf(self, doc: Documentation, template: str) -> bytes:
        """Export documentation to PDF format."""
        from reportlab.pdfgen import canvas
        from io import BytesIO
        
        buffer = BytesIO()
        pdf = canvas.Canvas(buffer)
        pdf.drawString(100, 800, doc.title)
        y_position = 750
        for block in doc.code_blocks:
            pdf.drawString(100, y_position, block.content)
            y_position -= 20
        pdf.save()
        
        return buffer.getvalue()

    def _export_docx(self, doc: Documentation, template: str) -> bytes:
        """Export documentation to DOCX format."""
        from docx import Document
        from io import BytesIO
        
        docx_doc = Document()
        docx_doc.add_heading(doc.title, 0)
        for block in doc.code_blocks:
            docx_doc.add_paragraph(block.content)
        
        buffer = BytesIO()
        docx_doc.save(buffer)
        return buffer.getvalue()

    def _markdown_to_html(self, markdown_content: str) -> str:
        """Convert markdown to HTML with basic formatting."""
        # Convert headers
        content = re.sub(r'^# (.*?)$', r'<h1>\1</h1>', markdown_content, flags=re.MULTILINE)
        content = re.sub(r'^## (.*?)$', r'<h2>\1</h2>', content, flags=re.MULTILINE)
        content = re.sub(r'^### (.*?)$', r'<h3>\1</h3>', content, flags=re.MULTILINE)
        
        # Convert code blocks
        content = re.sub(
            r'```(\w+)\n(.*?)\n```',
            lambda m: f'<div class="code-block"><div class="code-block-header">{m.group(1)}</div><pre><code>{m.group(2)}</code></pre></div>',
            content,
            flags=re.DOTALL
        )
        
        # Convert inline code
        content = re.sub(r'`(.*?)`', r'<code>\1</code>', content)
        
        # Convert bold and italic
        content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
        content = re.sub(r'\*(.*?)\*', r'<em>\1</em>', content)
        
        # Convert lists
        content = re.sub(r'^\- (.*?)$', r'<li>\1</li>', content, flags=re.MULTILINE)
        content = content.replace('<li>', '<ul><li>').replace('</li>\n\n', '</li></ul>\n\n')
        
        # Convert metrics JSON blocks
        content = re.sub(
            r'```json\n(.*?)\n```',
            lambda m: f'<div class="metrics">{self._format_metrics_html(m.group(1))}</div>',
            content,
            flags=re.DOTALL
        )
        
        # Convert paragraphs
        content = re.sub(r'\n\n(.*?)\n\n', r'\n\n<p>\1</p>\n\n', content)
        
        return content

    def _format_metrics_html(self, metrics_json: str) -> str:
        """Format metrics JSON as an HTML table."""
        try:
            metrics = json.loads(metrics_json)
            rows = []
            for key, value in metrics.items():
                # Convert snake_case to Title Case
                formatted_key = ' '.join(word.capitalize() for word in key.split('_'))
                if isinstance(value, (int, float)):
                    formatted_value = f"{value:.2f}" if isinstance(value, float) else str(value)
                else:
                    formatted_value = str(value)
                rows.append(f"<tr><th>{formatted_key}</th><td>{formatted_value}</td></tr>")
                
            return f"""
                <table class="metrics-table">
                    <tbody>
                        {''.join(rows)}
                    </tbody>
                </table>
            """
        except json.JSONDecodeError:
            return f"<pre>{metrics_json}</pre>"

    def _generate_toc(self, doc: Documentation) -> str:
        """Generate table of contents for detailed template."""
        toc_items = [
            "- [Overview](#overview)",
            "- [Code Blocks](#code-blocks)"
        ]
        
        for i, block in enumerate(doc.code_blocks, 1):
            first_line = block.content.split('\n')[0].strip()
            toc_items.append(f"  - [Block {i}: {first_line}](#block-at-line-{block.line_number})")
            
        if doc.metrics:
            toc_items.append("- [Project Metrics](#project-metrics)")
            
        return '\n'.join(toc_items)

    def _export_json(self, doc: Documentation, template: str) -> str:
        """Export documentation to JSON format."""
        return json.dumps({
            'title': doc.title,
            'description': doc.description,
            'language': doc.language,
            'metrics': doc.metrics,
            'code_blocks': [
                {
                    'content': block.content,
                    'language': block.language,
                    'line_number': block.line_number,
                    'metrics': {
                        'loc': len(block.content.split('\n')),
                        'complexity': self._calculate_complexity(block)
                    }
                }
                for block in doc.code_blocks
            ]
        }, indent=2)

    def _export_pdf(self, doc: Documentation, template: str) -> bytes:
        """Export documentation to PDF format using reportlab."""
        buffer = BytesIO()
        doc_pdf = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create content
        content = []
        
        # Add title
        content.append(Paragraph(f"# {doc.language} Documentation", styles['Heading1']))
        content.append(Spacer(1, 12))
        
        # Add code blocks with syntax highlighting
        for block in doc.code_blocks:
            content.append(Paragraph("Code:", styles['Heading2']))
            content.append(Preformatted(block.content, styles['Code']))
            content.append(Spacer(1, 12))
        
        doc_pdf.build(content)
        return buffer.getvalue()

    def _export_docx(self, doc: Documentation, template: str) -> bytes:
        """Export documentation to DOCX format."""
        docx_document = Document()
        docx_document.add_heading(f"{doc.language} Documentation", 0)
        
        for block in doc.code_blocks:
            docx_document.add_heading("Code:", level=2)
            docx_document.add_paragraph(block.content, style='Normal')
        
        buffer = BytesIO()
        docx_document.save(buffer)
        buffer.seek(0)
        return buffer.read()